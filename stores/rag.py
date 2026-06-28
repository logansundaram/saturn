import hashlib
import json
import re
import shutil
from collections import Counter
from pathlib import Path

from langchain_core.vectorstores import InMemoryVectorStore
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
import pypdf

from config import get_config
from core.llms import get_embeddings
from stores.document_registry import (
    manifest_entries,
    read_documents_manifest,
    register_rag_document,
    remove_rag_document,
)


# What the corpus ingests. Text formats load directly; PDFs are cleaned (furniture/hyphenation);
# HTML goes through trafilatura (already a dependency for web_extract) with a tag-strip fallback;
# CSV is prefixed with a column summary so chunks keep their schema; DOCX needs python-docx
# (requirements.txt) and fails that one file with a clear message when it's missing.
SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".html", ".htm", ".csv", ".docx"}


def documents_dir():
    """The corpus directory, resolved from config at call time so a live `/config paths.documents`
    change is honored without a restart."""
    return get_config().path("documents")


def iter_documents():
    """Yield the document files RAG will ingest — supported extensions under the corpus dir,
    recursively. The single source of truth for 'what counts as a document', shared by the
    ingest pipeline and the startup banner so their counts can't drift."""
    root = documents_dir()
    if not root.exists():
        return
    for file_path in root.glob("**/*"):
        # Skip dotfiles (notably `.manifest.md`, which document_registry writes *into* the corpus
        # dir) so the manifest isn't itself ingested as a corpus document.
        if file_path.name.startswith("."):
            continue
        if file_path.is_file() and file_path.suffix.lower() in SUPPORTED_EXTENSIONS:
            yield file_path


# ── on-disk cache (so the corpus isn't re-embedded every startup) ─────────────────────────────
# The vector store is dumped to `paths.cache/vectors.json` and the corpus fingerprint to
# `index.json` ({embedder, chunking, files: {source: {hash, chunk_ids}}}). On startup `sync()`
# reconciles the persisted store against what's on disk by content hash — embedding only
# new/changed files, dropping vectors for removed ones, loading the rest from the dump — so an
# unchanged corpus does zero embedding calls. The embedder id AND the chunking settings are part
# of the index: swapping embedders (a tier change) or editing `rag.chunk_size`/`chunk_overlap`
# invalidates every vector and forces a full rebuild, which keeps the cache from drifting out of
# sync with the model/settings that produced it.
_STORE_FILE = "vectors.json"
_INDEX_FILE = "index.json"


def _cache_dir() -> Path:
    d = get_config().path("cache")
    d.mkdir(parents=True, exist_ok=True)
    return d


def _store_path() -> Path:
    return _cache_dir() / _STORE_FILE


def _index_path() -> Path:
    return _cache_dir() / _INDEX_FILE


def _read_index() -> dict:
    p = _index_path()
    if p.exists():
        try:
            return json.loads(p.read_text(encoding="utf-8"))
        except Exception:
            pass
    return {"embedder": None, "files": {}}


def _write_index(index: dict) -> None:
    _index_path().write_text(json.dumps(index, indent=2), encoding="utf-8")


def _file_hash(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


# ── vector store (lazy, embedder-aware, disk-backed) ──────────────────────────────────────────
# Built/loaded on first use from the active tier's `embedder` (config.yaml). `sync()` owns the
# load-from-disk and incremental-embed logic; the bare `_build_store()` makes a fresh empty store
# for a full rebuild.
_embeddings = None
_vector_store = None
_store_embedder = None  # the embedder id the live store was built with


def _build_store() -> None:
    """Replace the live store with a fresh, empty one bound to the active embedder."""
    global _embeddings, _vector_store, _store_embedder
    _embeddings = get_embeddings()
    _vector_store = InMemoryVectorStore(_embeddings)
    _store_embedder = get_config().embedder_model


def _load_store() -> bool:
    """Load the live store from the on-disk dump, bound to the active embedder. Returns True on
    success; on failure (missing/corrupt dump) leaves a fresh empty store and returns False so the
    caller can fall back to a full rebuild rather than trust a stale index against an empty store."""
    global _embeddings, _vector_store, _store_embedder
    _embeddings = get_embeddings()
    _store_embedder = get_config().embedder_model
    try:
        _vector_store = InMemoryVectorStore.load(str(_store_path()), _embeddings)
        return True
    except Exception:
        _vector_store = InMemoryVectorStore(_embeddings)
        return False


def get_vector_store():
    """The active vector store. On first use it runs `sync()`, which loads the cached store from
    disk and reconciles it against the corpus — so callers never see an empty store just because
    nothing has triggered ingest yet this process."""
    if _vector_store is None:
        sync(verbose=False)
    return _vector_store


# ── document loading / chunking ───────────────────────────────────────────────────────────────
def _get_splitter() -> RecursiveCharacterTextSplitter:
    """The chunker, built from the `rag:` knobs in config.yaml at call time so a live `/config
    rag.chunk_size` edit is honored. The chunking params ride the index fingerprint (see
    `_chunking_fingerprint`), so a change forces a full re-embed — the stored vectors always
    match the live settings."""
    cfg = get_config()
    return RecursiveCharacterTextSplitter(
        chunk_size=int(cfg.get("rag.chunk_size", 1000)),
        chunk_overlap=int(cfg.get("rag.chunk_overlap", 150)),
    )


def _chunking_fingerprint() -> list:
    """The chunk-shaping settings as stored in index.json — when these differ from the persisted
    ones, sync() does a full rebuild (mirrors the embedder-change rule)."""
    cfg = get_config()
    return [int(cfg.get("rag.chunk_size", 1000)), int(cfg.get("rag.chunk_overlap", 150))]


def retrieval_k() -> int:
    """Chunks returned per knowledge-base search (`rag.k`). Read by the search_knowledge_base
    tool at call time, so a /config change applies immediately (no re-embed)."""
    return int(get_config().get("rag.k", 6))


# ── PDF text cleanup ──────────────────────────────────────────────────────────────────────────
# pypdf's raw extraction keeps page furniture (running headers/footers, page numbers) and hard
# line breaks mid-word. Both poison retrieval: furniture lines repeat into many chunks and match
# everything weakly; split words match nothing. Cleaned BEFORE chunking so the vectors only ever
# see content.
_DIGITS_RE = re.compile(r"\d+")


def _furniture_key(line: str) -> str:
    """Normalize a candidate header/footer line for cross-page comparison: digits collapse so
    'Page 3 of 12' and 'Page 4 of 12' read as the same repeated line."""
    return _DIGITS_RE.sub("#", line.strip().lower())


def _strip_repeated_furniture(pages: list[str], edge: int = 2, ratio: float = 0.6) -> list[str]:
    """Drop running headers/footers: a (digit-normalized) line that opens or closes most pages is
    page furniture, not content. Only the `edge` outermost lines of each page are candidates, so a
    sentence legitimately repeated mid-page is never touched. No-op for short documents (<3 pages),
    where 'repeated across pages' isn't meaningful."""
    if len(pages) < 3:
        return pages
    heads: Counter = Counter()
    feet: Counter = Counter()
    split_pages = [p.splitlines() for p in pages]
    edges = []  # per page: (head line indices, foot line indices) — the furniture candidates
    for lines in split_pages:
        content_idx = [i for i, ln in enumerate(lines) if ln.strip()]
        head_idx, foot_idx = content_idx[:edge], content_idx[-edge:]
        edges.append((head_idx, foot_idx))
        for i in head_idx:
            heads[_furniture_key(lines[i])] += 1
        for i in foot_idx:
            feet[_furniture_key(lines[i])] += 1
    threshold = max(3, int(len(pages) * ratio))
    head_junk = {k for k, c in heads.items() if c >= threshold}
    foot_junk = {k for k, c in feet.items() if c >= threshold}
    if not head_junk and not foot_junk:
        return pages

    cleaned = []
    for lines, (head_idx, foot_idx) in zip(split_pages, edges):
        drop = {i for i in head_idx if _furniture_key(lines[i]) in head_junk}
        drop |= {i for i in foot_idx if _furniture_key(lines[i]) in foot_junk}
        cleaned.append("\n".join(ln for i, ln in enumerate(lines) if i not in drop))
    return cleaned


def _normalize_pdf_text(text: str) -> str:
    """Repair extraction artifacts: rejoin words hyphenated across line breaks, strip trailing
    whitespace, collapse blank-line runs."""
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# Markdown heading levels that become section breadcrumbs on a chunk (deeper levels stay inline).
_MD_HEADERS = [("#", "h1"), ("##", "h2"), ("###", "h3")]


def _html_to_text(raw: str) -> str:
    """Readable text from an HTML corpus file. trafilatura (the same local extractor web_extract
    uses) does the heavy lifting; when it's unavailable or extracts nothing (a fragment, a
    minimal page), degrade to a crude tag-strip so the file still embeds as *something* readable
    rather than failing or embedding angle-bracket soup."""
    try:
        import trafilatura

        text = trafilatura.extract(raw) or ""
        if text.strip():
            return text
    except Exception:
        pass
    import html as _html

    text = re.sub(r"(?is)<(script|style)[^>]*>.*?</\1>", " ", raw)
    text = re.sub(r"(?s)<[^>]+>", " ", text)
    text = _html.unescape(text)
    text = re.sub(r"[ \t]+", " ", text)
    return re.sub(r"\n{3,}", "\n\n", text).strip()


def _csv_to_text(raw: str) -> str:
    """CSV prepared for chunking: a `[columns: …]` header line is prepended so every chunk —
    including ones far from the first row — can be traced back to the schema. The rows themselves
    stay verbatim (the chunker splits long files; values are what retrieval matches on)."""
    try:
        import csv as _csv
        import io

        first = next(_csv.reader(io.StringIO(raw)), None)
        if first and any(c.strip() for c in first):
            cols = ", ".join(c.strip() for c in first if c.strip())
            return f"[columns: {cols}]\n{raw}"
    except Exception:
        pass
    return raw


def _docx_to_text(path: Path) -> str:
    """Text from a .docx: paragraphs in order, plus table cells row by row (tab-joined) — the two
    places Word documents keep their prose. Needs python-docx; a missing package raises a clear,
    actionable error that sync() reports for THIS file while the rest of the corpus proceeds."""
    try:
        import docx  # python-docx
    except ImportError as exc:
        raise RuntimeError(
            "reading .docx needs the python-docx package — run `pip install python-docx`"
        ) from exc
    d = docx.Document(str(path))
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append("\t".join(cells))
    return "\n\n".join(parts)


def _load_file_docs(path: Path):
    """Load one corpus file into (source, [Document], full_text). PDFs become one cleaned Document
    per page (`page` in metadata; furniture stripped, hyphenation repaired); markdown is
    pre-sectioned by heading so chunks carry their section breadcrumb; HTML is extracted to
    readable text (trafilatura); CSV gets a schema header; DOCX extracts paragraphs + tables;
    plain text a single Document. `full_text` is what the manifest summarizer sees."""
    root = documents_dir()
    try:
        source = str(path.relative_to(root))
    except ValueError:
        # Screening a file BEFORE it is copied into the corpus (screen_file) — label by name.
        source = path.name
    suffix = path.suffix.lower()
    docs = []
    if suffix == ".pdf":
        reader = pypdf.PdfReader(str(path))
        raw_pages = [page.extract_text() or "" for page in reader.pages]
        page_texts = [_normalize_pdf_text(t) for t in _strip_repeated_furniture(raw_pages)]
        for page_num, text in enumerate(page_texts):
            if text.strip():
                docs.append(
                    Document(page_content=text, metadata={"source": source, "page": page_num + 1})
                )
        full_text = "\n\n".join(t for t in page_texts if t.strip())
    elif suffix == ".md":
        # errors="replace": one undecodable byte in a corpus file must degrade to a marker, not
        # raise — an uncaught UnicodeDecodeError here aborts sync() and breaks ALL retrieval
        # (get_vector_store runs sync on first use). Mirrors read_file / mentions.
        full_text = path.read_text(encoding="utf-8", errors="replace")
        docs = _markdown_sections(source, full_text)
    elif suffix in (".html", ".htm"):
        full_text = _html_to_text(path.read_text(encoding="utf-8", errors="replace"))
        docs.append(Document(page_content=full_text, metadata={"source": source}))
    elif suffix == ".csv":
        full_text = _csv_to_text(path.read_text(encoding="utf-8", errors="replace"))
        docs.append(Document(page_content=full_text, metadata={"source": source}))
    elif suffix == ".docx":
        full_text = _docx_to_text(path)
        docs.append(Document(page_content=full_text, metadata={"source": source}))
    else:
        full_text = path.read_text(encoding="utf-8", errors="replace")
        docs.append(Document(page_content=full_text, metadata={"source": source}))
    return source, docs, full_text


def _markdown_sections(source: str, full_text: str) -> list[Document]:
    """Split a markdown file into per-section Documents with h1/h2/h3 metadata, so each chunk can
    carry its heading breadcrumb (see `_chunks_for`). Falls back to one whole-file Document if the
    header splitter is unavailable or the file has no headings."""
    try:
        from langchain_text_splitters import MarkdownHeaderTextSplitter

        sections = MarkdownHeaderTextSplitter(_MD_HEADERS).split_text(full_text)
    except Exception:
        sections = []
    if not sections:
        return [Document(page_content=full_text, metadata={"source": source})]
    for d in sections:
        d.metadata["source"] = source
    return sections


def _chunks_for(source: str, docs):
    """Split a file's Documents into chunks with deterministic, source-scoped ids (`source::N`)
    so a file's vectors can be deleted/replaced wholesale on change or removal. A chunk from a
    markdown section gets its heading breadcrumb prepended (and kept in `section` metadata) — the
    embedder and the agent both see which part of the document the text came from."""
    chunks = _get_splitter().split_documents(docs)
    for c in chunks:
        crumb = " › ".join(
            v for v in (c.metadata.pop("h1", None), c.metadata.pop("h2", None),
                        c.metadata.pop("h3", None)) if v
        )
        if crumb:
            c.metadata["section"] = crumb
            c.page_content = f"§ {crumb}\n{c.page_content}"
    ids = [f"{source}::{i}" for i in range(len(chunks))]
    return chunks, ids


# ── corpus admission screening ───────────────────────────────────────────────────────────────
# A trojanized document is the quarantine side door: ingest it once and its payload re-presents
# on every RAG search. Retrieval is already quarantined (search_knowledge_base is untrusted), but
# the warning belongs at the SOURCE too: /docs add asks the human before admitting a flagged
# file; sync() reports what it admitted.


# Screen→ingest handoff: /docs add screens a file (full load + scan via screen_file), then
# ingest_file → sync() would load and scan the very same bytes again seconds later — for a large
# PDF that doubles the pypdf parse + full-text regex scan before the embed even starts. One slot,
# keyed by CONTENT HASH so it can never serve stale or wrong-file data (the corpus copy is
# byte-identical to the screened source), popped on use.
_SCREENED: dict = {}


def screen_file(src_path) -> list:
    """The instruction-shaped findings (quarantine.scan) in a would-be document's extracted
    text. Pure read — nothing is copied or embedded (the loaded docs + findings are cached by
    content hash for the ingest that typically follows; see _SCREENED). Returns [] when
    quarantine is off, or when the file can't be parsed (the loader will fail loudly at ingest;
    screening must never block what ingest would reject anyway)."""
    from trust import quarantine

    if not quarantine.active():
        return []
    path = Path(src_path).expanduser()
    try:
        h = _file_hash(path)
        _, docs, full_text = _load_file_docs(path)
    except Exception:
        return []
    findings = quarantine.scan(full_text)
    _SCREENED.clear()  # one slot — the handoff is immediate, never a cache to manage
    _SCREENED[h] = (docs, full_text, findings)
    return findings


def _admission_flags(full_text: str, findings: "list | None" = None) -> list[str]:
    """The flag KINDS in one admitted document ([] when quarantine is off) — sync() records
    these per file so the corpus never gains an injection payload silently. `findings` lets the
    screen-file handoff pass its already-computed scan instead of re-scanning the full text."""
    from trust import quarantine

    if not quarantine.active():
        return []
    if findings is None:
        findings = quarantine.scan(full_text)
    return sorted({f.kind for f in findings})


# ── sync: the single reconcile-against-disk entry point ─────────────────────────────────────────
def sync(*, force: bool = False, verbose: bool = True, on_file=None) -> dict:
    """Reconcile the persisted vector store + document manifest against the corpus on disk.

    Loads the cached store, then by content hash: embeds new/changed files, drops vectors +
    manifest entries for removed files, and leaves unchanged files alone. An embedder OR
    chunking-config change (or `force=True`, used by /docs sync --force) triggers a full
    re-embed. The manifest is additionally reconciled against disk directly, so a removed file
    loses its manifest entry even on a full rebuild (where the index was just reset and cannot
    name it). Re-dumps the store and rewrites the index only when something actually changed.
    Returns a stats dict: added / updated / removed / unchanged / rebuilt.

    `on_file(source, i, n)` (optional) is called before each file is embedded — the progress
    hook the /docs sync command renders, so a long re-embed isn't silent.

    A startup whose corpus hasn't changed does zero embedding calls — that's the whole point."""
    embedder = get_config().embedder_model
    chunking = _chunking_fingerprint()
    index = _read_index()
    full_rebuild = (
        force
        or index.get("embedder") != embedder
        or index.get("chunking") != chunking
        or not _store_path().exists()
    )

    if not full_rebuild and not _load_store():
        full_rebuild = True  # dump missing/corrupt: don't trust the index against an empty store

    if full_rebuild:
        _build_store()
        files: dict = {}
    else:
        files = dict(index.get("files", {}))

    store = _vector_store
    on_disk = {str(p.relative_to(documents_dir())): p for p in iter_documents()}
    stats = {
        "added": 0,
        "updated": 0,
        "removed": 0,
        "unchanged": 0,
        "rebuilt": full_rebuild,
        # (source, error) for files whose LOADER failed (corrupt docx, missing python-docx, …).
        # A bad file degrades to a reported skip instead of aborting the whole sync — one
        # unreadable document must never take down retrieval for the rest of the corpus.
        "failed": [],
        # (source, [kind, …]) for admitted files carrying instruction-shaped content — the
        # corpus admission warning (see screen_file above).
        "flagged": [],
    }

    # Files gone from disk: drop their vectors + manifest entry. `removed_sources` tracks the
    # index-level removals so the manifest reconcile below never double-counts one, and so the
    # dump/index rewrite at the end fires exactly when the store/index actually mutated.
    removed_sources: set = set()
    for source in [s for s in files if s not in on_disk]:
        ids = files[source].get("chunk_ids") or []
        if ids:
            store.delete(ids)
        remove_rag_document(source)
        del files[source]
        removed_sources.add(source)
        stats["removed"] += 1

    # Manifest orphans: entries for documents no longer on disk that the index walk above cannot
    # see. On a full rebuild `files` was just reset to {} (and a wiped cache/ dir — documented
    # "safe to delete" — has no index at all), so a deleted document's manifest block + cached
    # summary would otherwise survive forever: the fresh index only records on-disk files, so no
    # FUTURE sync would notice the orphan either. Reconcile the manifest against disk directly.
    # Vectors need no cleanup here — a rebuilt store starts empty, and the loop above already
    # deleted every indexed source's chunks. An orphan is, by definition, absent from `files`,
    # so this pass never touches the store or the index.
    for entry in manifest_entries(read_documents_manifest()):
        name = entry.get("name", "")
        if name and name not in on_disk and name not in removed_sources:
            remove_rag_document(name)
            stats["removed"] += 1

    # New or changed files: re-embed only those (hash-check first so the progress hook knows the
    # real total — unchanged files never count toward it).
    to_embed = []
    for source, path in on_disk.items():
        h = _file_hash(path)
        entry = files.get(source)
        if entry and entry.get("hash") == h:
            stats["unchanged"] += 1
            continue
        to_embed.append((source, path, h, entry))

    for i, (source, path, h, entry) in enumerate(to_embed, start=1):
        if on_file:
            on_file(source, i, len(to_embed))
        # Load BEFORE deleting the old vectors: if the loader fails (corrupt file, missing
        # optional package), the previously-embedded version stays searchable and the failure is
        # reported per-file instead of aborting the sync for the whole corpus. An EMBEDDING
        # failure (daemon down) still raises out as before — nothing can proceed without it.
        try:
            screened = _SCREENED.pop(h, None)
            if screened is not None:
                # The screen-file handoff: same bytes (content-hash keyed), already loaded and
                # scanned by /docs add moments ago. Relabel to the corpus source key — the
                # screen loaded from the ORIGINAL path, so its docs carry the bare basename.
                docs, full_text, findings = screened
                for d in docs:
                    d.metadata["source"] = source
            else:
                _src, docs, full_text = _load_file_docs(path)
                findings = None
            chunks, ids = _chunks_for(source, docs)
        except Exception as exc:
            stats["failed"].append((source, str(exc)))
            continue
        if entry and entry.get("chunk_ids"):
            store.delete(entry["chunk_ids"])  # replace the old vectors for a changed file
        if chunks:
            store.add_documents(chunks, ids=ids)
        register_rag_document(source, full_text)  # manifest summary (cached by hash downstream)
        files[source] = {"hash": h, "chunk_ids": ids}
        stats["updated" if entry else "added"] += 1
        kinds = _admission_flags(full_text, findings)
        if kinds:
            stats["flagged"].append((source, kinds))

    # Dump + index rewrite only when something mutated the store or the index. The common case —
    # a startup whose corpus hasn't changed — must not rewrite a multi-MB vectors.json on every
    # launch: besides the waste, each rewrite widens the crash window in which a kill mid-write
    # leaves a truncated dump that _load_store() treats as corrupt, silently forcing a full
    # re-embed of the whole corpus next run. Behavior-preserving: an unchanged store/index
    # round-trips to identical content anyway. Manifest-orphan cleanup above deliberately does
    # not trigger a rewrite (it changes neither the store nor `files`).
    if full_rebuild or to_embed or removed_sources:
        store.dump(str(_store_path()))
        _write_index({"embedder": embedder, "chunking": chunking, "files": files})

    if verbose:
        print(
            f"RAG cache synced ({documents_dir()}): "
            f"+{stats['added']} ~{stats['updated']} -{stats['removed']} "
            f"={stats['unchanged']}" + ("  [full rebuild]" if full_rebuild else "")
        )
        for source, err in stats["failed"]:
            print(f"  failed to load {source}: {err}")
        for source, kinds in stats["flagged"]:
            print(f"  ⚠ {source}: instruction-shaped content ({', '.join(kinds)}) — "
                  f"its search results are quarantined as untrusted")
    return stats


def sync_to_config() -> bool:
    """Re-embed the corpus if the configured embedder changed since the store was built. Returns
    True if it re-ingested. Call after a live model/tier change (e.g. /model, /config) so an
    embedder swap takes effect; a no-op when the embedder is unchanged."""
    # What embedder the cache was built with: the live store's if loaded, else the persisted
    # index's. Using the on-disk index (rather than forcing a rebuild whenever the store happens to
    # be unloaded) means an unrelated model rebind doesn't needlessly re-embed the whole corpus when
    # the embedder is actually unchanged.
    cached_embedder = _store_embedder if _vector_store is not None else _read_index().get("embedder")
    if cached_embedder == get_config().embedder_model:
        return False
    sync(force=True)
    return True


# ── corpus mutation (backs /docs add and /docs remove) ──────────────────────────────────────────
def ingest_file(src_path: str) -> dict:
    """Add a document to the corpus and embed it. If the file isn't already under the corpus dir
    it's copied in first. Returns the `sync()` stats."""
    p = Path(src_path).expanduser()
    if not p.exists():
        raise FileNotFoundError(src_path)
    if p.suffix.lower() not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"unsupported type '{p.suffix}'; supported: {sorted(SUPPORTED_EXTENSIONS)}"
        )
    root = documents_dir()
    root.mkdir(parents=True, exist_ok=True)
    dest = root / p.name
    if p.resolve() != dest.resolve():
        # Never silently clobber: the corpus keys documents by basename, so two different source
        # files sharing a name would otherwise overwrite each other and sync() would report the
        # data loss as an innocuous update. Identical bytes are a harmless re-add; anything else
        # requires the user to remove the existing document explicitly first.
        if dest.exists() and dest.read_bytes() != p.read_bytes():
            raise FileExistsError(
                f"a different document named '{p.name}' is already in the knowledge base — "
                f"if this replaces it, run `/docs remove {p.name}` first; otherwise rename "
                "the new file before adding it."
            )
        shutil.copy2(p, dest)
    return sync(verbose=False)


def forget_document(name: str) -> bool:
    """Remove a document from the corpus by relative source or basename. `sync()` then drops its
    vectors + manifest entry. Returns False if no matching file exists."""
    root = documents_dir()
    target = root / name
    if not target.exists():
        matches = [p for p in iter_documents() if p.name == name]
        if not matches:
            return False
        target = matches[0]
    target.unlink()
    sync(verbose=False)
    return True
